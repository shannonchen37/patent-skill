from __future__ import annotations

import re
from pathlib import Path

INTERNAL_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)
TRACE_LABEL = re.compile(r"\[I\d+-L\d+\]\s*")


def render_docx(source_dir: Path, output: Path) -> None:
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("Install the 'docx' extra: pip install -e '.[docx]'") from exc
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    tokens = {
        # Named override for this Chinese drafting artifact: Hiragino Sans GB
        # replaces Calibri so both Latin and CJK glyphs survive headless rendering.
        "Normal": ("Hiragino Sans GB", 11, "000000", 0, 6),
        "Heading 1": ("Hiragino Sans GB", 16, "2E74B5", 16, 8),
        "Heading 2": ("Hiragino Sans GB", 13, "2E74B5", 12, 6),
        "Heading 3": ("Hiragino Sans GB", 12, "1F4D78", 8, 4),
    }
    for style_name, (font_name, size, color, before, after) in tokens.items():
        style = document.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.element.rPr.rFonts.set(qn("w:ascii"), font_name)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.element.rPr.rFonts.set(qn("w:cs"), font_name)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
    if "Patent Claim" not in document.styles:
        claim_style = document.styles.add_style("Patent Claim", WD_STYLE_TYPE.PARAGRAPH)
        claim_style.base_style = document.styles["Normal"]
        claim_style.paragraph_format.keep_together = True
    for filename, title in (
        ("claims-v2.md", "权利要求书"),
        ("specification.md", "说明书"),
        ("abstract.md", "说明书摘要"),
        ("figures.md", "附图说明"),
    ):
        path = source_dir / filename
        if not path.exists():
            continue
        document.add_heading(title, level=1)
        text = INTERNAL_COMMENT.sub("", path.read_text(encoding="utf-8"))
        if filename == "claims-v2.md":
            text = TRACE_LABEL.sub("", text)
        for line in text.splitlines():
            if line.startswith("# "):
                document.add_heading(line[2:], level=2)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=3)
            elif line.strip():
                style = "Patent Claim" if filename == "claims-v2.md" else "Normal"
                document.add_paragraph(line.strip(), style=style)
    # LibreOffice on macOS may ignore East Asian font settings inherited only
    # from paragraph styles, so reinforce the named CJK override on each run.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Hiragino Sans GB"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Hiragino Sans GB")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
            run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Hiragino Sans GB")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
